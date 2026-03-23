"""
文档解析服务 — 上传文件 → 文本提取 → 章节切分 → 入库 → 向量化

支持格式: .docx（python-docx 直接读取）、.doc（textutil 转换）
解析流程: 文件 → 纯文本 → 按章/节/条切分 → std_document + std_clause

注意: 此服务作为 Web 请求中的同步/异步调用，非 CLI 脚本。
"""
import os
import re
import subprocess
import tempfile
from typing import Optional

from docx import Document as DocxDocument
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.standard import StdDocument, StdClause


# ========== 文本提取 ==========

def extract_text_from_docx(file_path: str) -> str:
    """从 .docx 文件提取纯文本"""
    doc = DocxDocument(file_path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def extract_text_from_doc(file_path: str) -> str:
    """从 .doc 文件提取纯文本（macOS textutil）"""
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        subprocess.run(
            ["textutil", "-convert", "txt", "-output", tmp_path, file_path],
            check=True, capture_output=True
        )
        with open(tmp_path, "r", encoding="utf-8") as f:
            return f.read()
    except subprocess.CalledProcessError:
        return ""
    finally:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


def extract_text(file_path: str, filename: str) -> str:
    """根据文件扩展名选择提取方式"""
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext == ".doc":
        return extract_text_from_doc(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    else:
        raise ValueError(f"不支持的文件格式: {ext}，仅支持 .doc/.docx/.txt")


# ========== 章节切分 ==========

# 匹配中文章节结构的正则
RE_CHAPTER = re.compile(r'^第([一二三四五六七八九十百]+)章\s*(.*)')
RE_SECTION = re.compile(r'^第([一二三四五六七八九十百]+)节\s*(.*)')
RE_ITEM = re.compile(r'^([一二三四五六七八九十]+)、\s*(.*)')


def _sanitize(s: str) -> str:
    """清洗文本：移除 NUL 和控制字符"""
    s = s.replace('\x00', '')
    s = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f]', '', s)
    return s.strip()


def parse_to_clauses(text: str, doc_title: str) -> list[dict]:
    """
    按章/节/条切分文本为条款列表

    返回格式: [{"clause_no": "第一章", "title": "...", "content": "...", "level": 1}, ...]
    level: 1=章, 2=节, 3=条目
    """
    lines = text.splitlines()
    clauses: list[dict] = []
    cur_chapter = ""
    cur_section = ""
    cur_clause: dict | None = None

    def flush():
        nonlocal cur_clause
        if cur_clause and len(cur_clause["content"].strip()) > 10:
            cur_clause["content"] = _sanitize(cur_clause["content"])
            cur_clause["title"] = _sanitize(cur_clause["title"])
            clauses.append(cur_clause)
        cur_clause = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 章级
        m = RE_CHAPTER.match(line)
        if m:
            flush()
            cur_chapter = f"第{m.group(1)}章 {m.group(2).strip()}"
            cur_section = ""
            cur_clause = {
                "clause_no": f"第{m.group(1)}章",
                "title": cur_chapter,
                "content": "",
                "level": 1,
            }
            continue

        # 节级
        m = RE_SECTION.match(line)
        if m:
            flush()
            cur_section = f"第{m.group(1)}节 {m.group(2).strip()}"
            path = f"{cur_chapter} > {cur_section}" if cur_chapter else cur_section
            cur_clause = {
                "clause_no": f"第{m.group(1)}节",
                "title": path,
                "content": "",
                "level": 2,
            }
            continue

        # 条目级
        m = RE_ITEM.match(line)
        if m:
            flush()
            item_title = m.group(2).strip()
            parts = [cur_chapter, cur_section, item_title]
            path = " > ".join(p for p in parts if p)
            cur_clause = {
                "clause_no": m.group(1),
                "title": path,
                "content": line + "\n",
                "level": 3,
            }
            continue

        # 普通行 → 追加到当前条款
        if cur_clause:
            cur_clause["content"] += line + "\n"
        else:
            # 文档开头没有章节标记的内容 → 创建一个"前言"条款
            cur_clause = {
                "clause_no": "前言",
                "title": f"{doc_title} > 前言",
                "content": line + "\n",
                "level": 0,
            }

    flush()
    return clauses


# ========== 入库服务 ==========

class DocumentParserService:
    """文档解析入库服务 — 从上传文件到数据库的完整 pipeline"""

    def __init__(self, session: AsyncSession):
        self.session = session

    async def parse_and_ingest(
        self,
        file_path: str,
        filename: str,
        doc_type: str = "安全规程",
        version: str = "v1.0",
        tenant_id: int = 1,
        created_by: int = 1,
    ) -> dict:
        """
        完整流程: 文件 → 文本提取 → 章节切分 → 入库

        返回: {"document_id": int, "title": str, "clause_count": int}
        """
        # 1. 提取文本
        text = extract_text(file_path, filename)
        if not text or len(text.strip()) < 20:
            raise ValueError("文件内容为空或过短，无法解析")

        # 2. 推导文档标题（去掉扩展名和路径）
        title = os.path.splitext(filename)[0]

        # 3. 创建文档记录
        doc = StdDocument(
            title=title,
            doc_type=doc_type,
            version=version,
            is_current=True,
            tenant_id=tenant_id,
            created_by=created_by,
        )
        self.session.add(doc)
        await self.session.flush()  # 获取 doc.id

        # 4. 章节切分
        clause_list = parse_to_clauses(text, title)

        # 5. 批量写入条款
        for c in clause_list:
            clause = StdClause(
                document_id=doc.id,
                clause_no=c["clause_no"],
                title=c["title"],
                content=c["content"],
                level=c["level"],
            )
            self.session.add(clause)

        await self.session.commit()

        return {
            "document_id": doc.id,
            "title": title,
            "doc_type": doc_type,
            "clause_count": len(clause_list),
        }

    async def vectorize_document(self, doc_id: int) -> int:
        """
        对指定文档的所有条款生成 Embedding 向量

        返回: 向量化的条款数
        """
        try:
            from google import genai

            api_key = os.getenv("GEMINI_API_KEY", "")
            if not api_key:
                return 0

            client = genai.Client(api_key=api_key)

            # 查询该文档所有未向量化的条款
            from sqlalchemy import select, and_
            stmt = select(StdClause).where(
                and_(
                    StdClause.document_id == doc_id,
                    StdClause.embedding.is_(None),
                )
            )
            result = await self.session.execute(stmt)
            clauses = result.scalars().all()

            if not clauses:
                return 0

            count = 0
            batch_size = 20

            for i in range(0, len(clauses), batch_size):
                batch = clauses[i:i + batch_size]
                texts = []
                for c in batch:
                    # 拼接 标题 + 条款号 + 内容 作为 embedding 输入
                    t = f"{c.title} {c.clause_no} {c.content or ''}"
                    texts.append(t[:2000])  # 截断保护

                # 调用 Gemini Embedding API
                response = client.models.embed_content(
                    model="gemini-embedding-001",
                    contents=texts,
                    config={"output_dimensionality": 1536},
                )

                # 更新数据库
                for j, emb in enumerate(response.embeddings):
                    batch[j].embedding = emb.values

                await self.session.commit()
                count += len(batch)

            return count

        except Exception as e:
            print(f"⚠️ 向量化失败: {e}")
            return 0
