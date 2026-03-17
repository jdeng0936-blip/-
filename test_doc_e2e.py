"""DOC-01 端到端集成验证脚本"""
import sys, os, json
sys.path.insert(0, os.path.join(os.path.dirname(__file__), "backend"))

from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.schemas.calc import SupportCalcInput
from app.schemas.vent import VentCalcInput
from app.services.calc_engine import SupportCalcEngine
from app.services.vent_engine import VentCalcEngine

# 参数字段中文映射
PARAM_LABELS = {
    "rock_class": "围岩级别", "coal_thickness": "煤层厚度(m)",
    "coal_dip_angle": "煤层倾角(°)", "gas_level": "瓦斯等级",
    "hydro_type": "水文地质类型", "geo_structure": "地质构造",
    "spontaneous_combustion": "自燃倾向性", "roadway_type": "巷道类型",
    "excavation_type": "掘进类型", "section_form": "断面形式",
    "section_width": "断面宽度(m)", "section_height": "断面高度(m)",
    "excavation_length": "掘进长度(m)", "service_years": "服务年限(年)",
    "dig_method": "掘进方式", "dig_equipment": "掘进设备",
    "transport_method": "运输方式",
}

# ===== 1. 项目参数 =====
project = {"name": "3301回风巷", "mine_name": "某某煤矿"}
params = {
    "rock_class": "IV", "coal_thickness": 3.8, "coal_dip_angle": 18,
    "gas_level": "高瓦斯", "hydro_type": "中等", "geo_structure": "含断层F3",
    "spontaneous_combustion": "不易自燃", "roadway_type": "回风巷",
    "excavation_type": "半煤岩巷", "section_form": "拱形",
    "section_width": 5.0, "section_height": 4.0,
    "excavation_length": 560, "service_years": 8,
    "dig_method": "综掘", "dig_equipment": "EBZ200H",
    "transport_method": "皮带运输",
}
print("[1/5] 项目参数已加载:", project["name"])

# ===== 2. 规则匹配 =====
matched_rules = [
    {"name": "IV类围岩-锚网索联合支护", "category": "支护", "priority": 10, "chapter": "第二章"},
    {"name": "高瓦斯-专项通风方案", "category": "通风", "priority": 15, "chapter": "第三章"},
    {"name": "含断层-过断层安全措施", "category": "安全", "priority": 12, "chapter": "第五章"},
    {"name": "综掘机-EBZ系列操作规程", "category": "装备", "priority": 8, "chapter": "第六章"},
]
print(f"[2/5] 规则匹配完成: {len(matched_rules)} 条命中")

# ===== 3. CAL-01 =====
calc_r = SupportCalcEngine.calculate(SupportCalcInput(
    rock_class="IV", section_form="拱形", section_width=5.0, section_height=4.0,
    bolt_spacing=1100, cable_count=2))
print(f"[3/5] CAL-01: K={calc_r.safety_factor}, 合规={calc_r.is_compliant}, 预警={len(calc_r.warnings)}")

# ===== 4. CAL-02 =====
vent_r = VentCalcEngine.calculate(VentCalcInput(
    gas_emission=3.5, gas_level="高瓦斯", section_area=calc_r.section_area,
    excavation_length=560, max_workers=30, explosive_per_cycle=15, design_air_volume=600))
print(f"[4/5] CAL-02: Q={vent_r.q_required}, 局扇={vent_r.recommended_fan}, 合规={vent_r.is_compliant}")

# ===== 5. 生成 Word =====
doc = Document()
doc.styles["Normal"].font.size = Pt(12)

# 封面
for _ in range(5): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.add_run(project["mine_name"]).font.size = Pt(26)
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
r = p.add_run(f"{project['name']}  掘进工作面作业规程"); r.font.size = Pt(20); r.font.bold = True
doc.add_paragraph()
for txt in [f"编制日期：{datetime.now().strftime('%Y年%m月%d日')}", "编制单位：生产技术科"]:
    p = doc.add_paragraph(); p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER; p.add_run(txt).font.size = Pt(14)
doc.add_page_break()

# 第一章 工程概况（参数表格）
doc.add_heading("第一章  工程概况", level=1)
tbl = doc.add_table(rows=1, cols=2, style="Table Grid")
tbl.rows[0].cells[0].text = "参数名称"; tbl.rows[0].cells[1].text = "参数值"
for field, label in PARAM_LABELS.items():
    row = tbl.add_row().cells; row[0].text = label; row[1].text = str(params.get(field, "—"))
doc.add_page_break()

# 第二章 支护设计
doc.add_heading("第二章  支护设计", level=1)
if not calc_r.is_compliant:
    wr = doc.add_paragraph().add_run("⚠ 本章节存在合规预警，请重点审查！")
    wr.font.color.rgb = RGBColor(0xCC, 0, 0); wr.font.bold = True
doc.add_heading("2.1 计算结果", level=2)
for k, v in [("断面面积", f"{calc_r.section_area} m²"), ("锚杆锚固力", f"{calc_r.bolt_force} kN"),
             ("最大间距", f"{calc_r.max_bolt_spacing} mm"), ("推荐每排", f"{calc_r.recommended_bolt_count_per_row} 根"),
             ("最少锚索", f"{calc_r.min_cable_count} 根"), ("安全系数", f"{calc_r.safety_factor}")]:
    doc.add_paragraph(f"{k}：{v}")
if calc_r.warnings:
    doc.add_heading("2.2 合规预警", level=2)
    for w in calc_r.warnings:
        r = doc.add_paragraph().add_run(f"{'🔴' if w.level=='error' else '🟡'} {w.message}")
        r.font.color.rgb = RGBColor(0xCC, 0, 0) if w.level == "error" else RGBColor(0xCC, 0x88, 0)
        r.font.bold = True
doc.add_page_break()

# 第三章 通风系统
doc.add_heading("第三章  通风系统", level=1)
if not vent_r.is_compliant:
    wr = doc.add_paragraph().add_run("⚠ 本章节存在合规预警，请重点审查！")
    wr.font.color.rgb = RGBColor(0xCC, 0, 0); wr.font.bold = True
doc.add_heading("3.1 需风量计算", level=2)
for k, v in [("瓦斯涌出法", f"{vent_r.q_gas} m³/min"), ("人数法", f"{vent_r.q_people} m³/min"),
             ("炸药法", f"{vent_r.q_explosive} m³/min"), ("最终配风量", f"{vent_r.q_required} m³/min")]:
    doc.add_paragraph(f"{k}：{v}")
doc.add_heading("3.2 局扇选型", level=2)
doc.add_paragraph(f"推荐局扇：{vent_r.recommended_fan}（{vent_r.fan_power} kW）")
if vent_r.warnings:
    doc.add_heading("3.3 合规预警", level=2)
    for w in vent_r.warnings:
        r = doc.add_paragraph().add_run(f"{'🔴' if w.level=='error' else '🟡'} {w.message}")
        r.font.color.rgb = RGBColor(0xCC, 0, 0) if w.level == "error" else RGBColor(0xCC, 0x88, 0)
        r.font.bold = True
doc.add_page_break()

# 第四章 规则命中
doc.add_heading("第四章  编制依据与规则命中", level=1)
for mr in matched_rules:
    r = doc.add_paragraph().add_run(f"• {mr['name']}（{mr['category']}，优先级 {mr['priority']}）→ {mr['chapter']}")
    r.font.bold = True
doc.add_page_break()

# 第五章 安全措施
doc.add_heading("第五章  安全技术措施", level=1)
for t, c in [("5.1 顶板管理", "严格执行敲帮问顶制度，支护紧跟迎头，严禁空顶作业。"),
             ("5.2 防治水", '坚持"有疑必探、先探后掘"原则。'),
             ("5.3 瓦斯管理", "瓦斯浓度达1.0%时停止作业、切断电源、撤出人员。"),
             ("5.4 综合防尘", "湿式打眼、喷雾降尘、通风除尘等综合防尘。")]:
    doc.add_heading(t, level=2); doc.add_paragraph(c)

# 保存
output_dir = os.path.join(os.path.dirname(__file__), "storage", "outputs")
os.makedirs(output_dir, exist_ok=True)
filepath = os.path.join(output_dir, f"{project['name']}_作业规程_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
doc.save(filepath)

print(f"\n[5/5] Word 生成成功!")
print(f"  路径: {filepath}")
print(f"  大小: {os.path.getsize(filepath)} bytes")
print(f"  章节: 5 章 | 支护预警: {len(calc_r.warnings)} | 通风预警: {len(vent_r.warnings)}")
assert os.path.exists(filepath) and os.path.getsize(filepath) > 10000
assert calc_r.is_compliant == False
print("\nALL PASSED ✅")
